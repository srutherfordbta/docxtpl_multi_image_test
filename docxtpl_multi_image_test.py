import docx
import docxtpl

TEMPLATE_FILE = 'images_template.docx'
OUT_FILE_BEFORE = 'images_before.docx'
OUT_FILE_AFTER = 'images_after.docx'
FOOTER = 'footer.png'

doc = docxtpl.DocxTemplate(TEMPLATE_FILE)
context = {'images': [docxtpl.InlineImage(doc, 'image_a.png'),
                      docxtpl.InlineImage(doc, 'image_b.png')]}
doc.render(context)
doc.save(OUT_FILE_BEFORE)

# Add Footer to the Document
document = docx.Document(OUT_FILE_BEFORE)
section = document.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0]
# Adjust Identations to overcome margins
paragraph_format = paragraph.paragraph_format
# -0.5 to override the margin settings
paragraph_format.left_indent = docx.shared.Inches(-0.5)
paragraph_format.right_indent = docx.shared.Inches(-0.5)
run = paragraph.add_run()
# Add the footer image
para_image = run.add_picture(FOOTER, width=docx.shared.Inches(8.5))
# Add the hyperlink
r_id = paragraph.part.relate_to('http://mywebsite.com',
                                docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                                is_external=True)
hyperlink = docx.oxml.shared.OxmlElement('a:hlinkClick')
hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
para_image._inline.docPr.append(hyperlink)
document.save(OUT_FILE_AFTER)
